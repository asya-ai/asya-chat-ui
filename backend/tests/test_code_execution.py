from __future__ import annotations

import base64
import os
import stat
from uuid import UUID, uuid4

from app.models import AgentSource, AgentSourceKind, AgentSourceStatus, ChatMessageAttachment
from app.models.entities import ChatCoworkDocument, CoworkFormat
from app.services.tools.code_execution import (
    _EXEC_BOOTSTRAP,
    _collect_outputs,
    _cowork_module_names,
    _sync_cowork_workspace,
    _validate_imports,
    _write_cowork_workspace,
    _write_inputs,
    _write_project_inputs,
    cowork_exec_path,
    project_source_exec_path,
)


def _attachment(*, attachment_id: str, file_name: str, payload: bytes) -> ChatMessageAttachment:
    return ChatMessageAttachment(
        id=UUID(attachment_id),
        message_id=uuid4(),
        file_name=file_name,
        content_type="application/octet-stream",
        data_base64=base64.b64encode(payload).decode("ascii"),
    )


def _source(
    *,
    source_id: str,
    title: str,
    content_text: str,
    file_name: str | None = None,
    file_path: str | None = None,
    content_type: str | None = "text/plain",
) -> AgentSource:
    return AgentSource(
        id=UUID(source_id),
        agent_id=uuid4(),
        kind=AgentSourceKind.file if file_name else AgentSourceKind.text,
        title=title,
        file_name=file_name,
        content_type=content_type,
        content_text=content_text,
        file_path=file_path,
        status=AgentSourceStatus.ready,
    )


def test_validate_imports_allows_presentation_dependencies() -> None:
    _validate_imports("import geopandas as gpd\nimport shapefile\nfrom pptx import Presentation")


def test_validate_imports_allows_scientific_and_document_stack() -> None:
    _validate_imports(
        "\n".join(
            [
                "import numpy",
                "import scipy",
                "import pandas",
                "import matplotlib",
                "import sklearn",
                "import statsmodels",
                "from PIL import Image",
                "import cv2",
                "import openpyxl",
                "import docx",
                "import pptx",
                "import reportlab",
                "import odf",
                "import pypandoc",
                "import bs4",
                "import sympy",
                "import shapely",
                "import networkx",
                "import xarray",
                "import polars",
                "import torch",
                "import tensorflow",
                "import qiskit",
                "import rdkit",
                "import math",
                "import statistics",
                "import json",
                "import csv",
                "import re",
                "import pathlib",
                "import datetime",
                "import sqlite3",
            ]
        )
    )


def test_validate_imports_blocks_network_and_host_introspection() -> None:
    for code in (
        "import requests",
        "import httpx",
        "import psutil",
        "from fake_useragent import UserAgent",
    ):
        try:
            _validate_imports(code)
        except ValueError as exc:
            assert "Import not allowed" in str(exc)
        else:
            raise AssertionError(f"Expected import rejection for: {code}")


def test_validate_imports_blocks_subprocess() -> None:
    for code in (
        "import subprocess",
        "import subprocess as sp",
        "from subprocess import run",
        "from subprocess import Popen, PIPE",
    ):
        try:
            _validate_imports(code)
        except ValueError as exc:
            assert "subprocess" in str(exc).lower()
        else:
            raise AssertionError(f"Expected import rejection for: {code}")
    try:
        _validate_imports("import subprocess", extra_allowed={"subprocess"})
    except ValueError as exc:
        assert "subprocess" in str(exc).lower()
    else:
        raise AssertionError("subprocess must stay blocked even if extra_allowed")


def test_validate_imports_allows_cowork_modules() -> None:
    _validate_imports(
        "import process_trs\nfrom process_trs import process_file_pair\nimport pandas as pd",
        extra_allowed={"process_trs"},
    )


def test_validate_imports_rejects_unknown_local_module() -> None:
    try:
        _validate_imports("import process_trs")
    except ValueError as exc:
        assert "Import not allowed" in str(exc)
    else:
        raise AssertionError("Expected import rejection for process_trs")


def test_cowork_module_names_from_py_files() -> None:
    py_doc = ChatCoworkDocument(
        chat_id=uuid4(),
        title="TRS",
        file_name="process_trs.py",
        format=CoworkFormat.code,
        content="def process_file_pair():\n    return True\n",
    )
    csv_doc = ChatCoworkDocument(
        chat_id=uuid4(),
        title="Sales",
        file_name="sales.csv",
        format=CoworkFormat.csv,
        content="a,b\n1,2\n",
    )
    assert _cowork_module_names([py_doc, csv_doc]) == {"process_trs"}


def test_exec_bootstrap_blocks_subprocess_and_adds_cowork_path(tmp_path) -> None:
    import subprocess as host_subprocess
    import sys

    cowork_dir = tmp_path / "cowork"
    cowork_dir.mkdir()
    (cowork_dir / "process_trs.py").write_text(
        "import subprocess\nVALUE = 42\n",
        encoding="utf-8",
    )
    (tmp_path / "main.py").write_text(
        "import process_trs\n"
        "print('cowork', process_trs.VALUE)\n"
        "try:\n"
        "    import subprocess\n"
        "    subprocess.run(['true'])\n"
        "except RuntimeError:\n"
        "    print('blocked')\n"
        "else:\n"
        "    raise SystemExit('subprocess.run succeeded')\n",
        encoding="utf-8",
    )
    bootstrap = (
        _EXEC_BOOTSTRAP.replace("/workspace/cowork", str(cowork_dir)).replace(
            "/workspace/main.py", str(tmp_path / "main.py")
        )
    )
    (tmp_path / "_sandbox_bootstrap.py").write_text(bootstrap, encoding="utf-8")
    result = host_subprocess.run(
        [sys.executable, str(tmp_path / "_sandbox_bootstrap.py")],
        cwd=tmp_path,
        capture_output=True,
        text=True,
        check=False,
    )
    assert result.returncode == 0, result.stderr
    assert "cowork 42" in result.stdout
    assert "blocked" in result.stdout


def test_write_inputs_stages_raw_attachments_without_preprocessing(tmp_path) -> None:
    shp_id = "11111111-1111-1111-1111-111111111111"
    zip_id = "22222222-2222-2222-2222-222222222222"
    attachments = [
        _attachment(attachment_id=shp_id, file_name="roads.shp", payload=b"shp-bytes"),
        _attachment(attachment_id=zip_id, file_name="bundle.zip", payload=b"zip-bytes"),
    ]

    inputs = _write_inputs(attachments, tmp_path)

    assert inputs == [
        {
            "name": "roads.shp",
            "path": f"/inputs/{shp_id}_roads.shp",
            "content_type": "application/octet-stream",
        },
        {
            "name": "bundle.zip",
            "path": f"/inputs/{zip_id}_bundle.zip",
            "content_type": "application/octet-stream",
        },
    ]
    assert (tmp_path / f"{shp_id}_roads.shp").read_bytes() == b"shp-bytes"
    assert (tmp_path / f"{zip_id}_bundle.zip").read_bytes() == b"zip-bytes"


def test_write_project_inputs_stages_content_text_under_project_dir(tmp_path) -> None:
    source_id = "33333333-3333-3333-3333-333333333333"
    source = _source(
        source_id=source_id,
        title="Sales Report",
        content_text="region,amount\nEU,10\n",
        file_name="sales.csv",
        content_type="text/csv",
    )

    inputs = _write_project_inputs([source], tmp_path)

    assert inputs == [
        {
            "name": "sales.csv",
            "path": f"/inputs/project/{source_id}_sales.csv",
            "content_type": "text/csv",
            "source_id": source_id,
            "title": "Sales Report",
        }
    ]
    assert (tmp_path / "project" / f"{source_id}_sales.csv").read_text(
        encoding="utf-8"
    ) == "region,amount\nEU,10\n"
    assert project_source_exec_path(source) == f"/inputs/project/{source_id}_sales.csv"


def test_write_project_inputs_prefers_stored_file_bytes(tmp_path, monkeypatch) -> None:
    source_id = "44444444-4444-4444-4444-444444444444"
    stored = tmp_path / "original.xlsx"
    stored.write_bytes(b"xlsx-bytes")
    source = _source(
        source_id=source_id,
        title="Workbook",
        content_text="extracted text only",
        file_name="workbook.xlsx",
        file_path="agents/x/sources/workbook.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    monkeypatch.setattr(
        "app.services.tools.code_execution.maybe_read_file_bytes",
        lambda path: stored.read_bytes() if path else None,
    )

    inputs = _write_project_inputs([source], tmp_path)

    assert inputs[0]["path"] == f"/inputs/project/{source_id}_workbook.xlsx"
    assert (tmp_path / "project" / f"{source_id}_workbook.xlsx").read_bytes() == b"xlsx-bytes"


def test_collect_outputs_keeps_large_files_as_attachments(monkeypatch, tmp_path) -> None:
    monkeypatch.setattr(
        "app.services.tools.code_execution.settings.exec_max_output_file_bytes",
        3,
    )
    monkeypatch.setattr(
        "app.services.tools.code_execution.settings.attachments_max_file_bytes",
        20,
    )
    monkeypatch.setattr(
        "app.services.tools.code_execution.settings.attachments_max_total_bytes",
        20,
    )
    (tmp_path / "image.png").write_bytes(b"larger-than-inline")

    attachments, output_items = _collect_outputs(tmp_path)

    assert attachments == [
        {
            "file_name": "image.png",
            "content_type": "image/png",
            "data_base64": base64.b64encode(b"larger-than-inline").decode("ascii"),
        }
    ]
    assert output_items == []


def test_collect_outputs_rejects_symlinks(tmp_path) -> None:
    target = tmp_path.parent / f"secret-{tmp_path.name}.txt"
    target.write_text("should-not-leak", encoding="utf-8")
    try:
        link = tmp_path / "symlink-collector-test.txt"
        link.symlink_to(target)
        (tmp_path / "ok.txt").write_bytes(b"safe")

        attachments, output_items = _collect_outputs(tmp_path)

        assert [item["file_name"] for item in attachments] == ["ok.txt"]
        assert attachments[0]["data_base64"] == base64.b64encode(b"safe").decode("ascii")
        assert [item["file_name"] for item in output_items] == ["ok.txt"]
    finally:
        target.unlink(missing_ok=True)


def test_collect_outputs_rejects_symlink_to_outside(tmp_path) -> None:
    outside = tmp_path.parent / f"outside-{tmp_path.name}.txt"
    outside.write_text("host-secret", encoding="utf-8")
    try:
        link = tmp_path / "escape.txt"
        link.symlink_to(outside)

        attachments, output_items = _collect_outputs(tmp_path)

        assert attachments == []
        assert output_items == []
    finally:
        outside.unlink(missing_ok=True)


def test_collect_outputs_rejects_fifo(tmp_path) -> None:
    fifo = tmp_path / "pipe.fifo"
    os.mkfifo(fifo)
    assert stat.S_ISFIFO(fifo.lstat().st_mode)
    (tmp_path / "ok.txt").write_bytes(b"ok")

    attachments, _output_items = _collect_outputs(tmp_path)

    assert [item["file_name"] for item in attachments] == ["ok.txt"]


def test_write_and_sync_cowork_workspace(tmp_path) -> None:
    from sqlmodel import Session, SQLModel, create_engine
    from sqlalchemy.pool import StaticPool

    from app.models.entities import Chat, Org, User

    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        org = Org(name="Org")
        session.add(org)
        session.commit()
        session.refresh(org)
        user = User(email="code-cowork@example.com", hashed_password="x")
        session.add(user)
        session.commit()
        session.refresh(user)
        chat = Chat(org_id=org.id, user_id=user.id, title="C")
        session.add(chat)
        session.commit()
        session.refresh(chat)

        doc = ChatCoworkDocument(
            chat_id=chat.id,
            title="Sales",
            file_name="sales.csv",
            format=CoworkFormat.csv,
            content="a,b\n1,2\n",
            version=1,
            last_assistant_version=1,
            content_at_assistant_version="a,b\n1,2\n",
            is_active=True,
        )
        session.add(doc)
        session.commit()
        session.refresh(doc)

        listing, path_to_id, snapshots = _write_cowork_workspace([doc], tmp_path)
        assert listing[0]["path"] == "/workspace/cowork/sales.csv"
        assert cowork_exec_path(doc) == "/workspace/cowork/sales.csv"
        assert (tmp_path / "cowork" / "sales.csv").read_text(encoding="utf-8") == "a,b\n1,2\n"
        assert (tmp_path / "cowork" / "manifest.json").is_file()

        # Unchanged content → no sync
        assert (
            _sync_cowork_workspace(
                session,
                chat.id,
                tmp_path,
                path_to_id=path_to_id,
                snapshots=snapshots,
            )
            == []
        )

        (tmp_path / "cowork" / "sales.csv").write_text("a,b\n1,2\n3,4\n", encoding="utf-8")
        updated = _sync_cowork_workspace(
            session,
            chat.id,
            tmp_path,
            path_to_id=path_to_id,
            snapshots=snapshots,
        )
        assert len(updated) == 1
        assert updated[0]["synced_path"] == "/workspace/cowork/sales.csv"
        assert updated[0]["content"] == "a,b\n1,2\n3,4\n"
        assert updated[0]["version"] == 2
        session.refresh(doc)
        assert doc.content == "a,b\n1,2\n3,4\n"
        assert doc.version == 2
